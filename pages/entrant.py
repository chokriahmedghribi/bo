import streamlit as st
from datetime import datetime
import os
from sqlalchemy.orm import Session
from database.models import init_db, Entrant, Contact
from database.crud import create_entrant, get_entrants, get_contacts, create_contact
import pandas as pd

st.set_page_config(page_title="البريد الوارد", layout="wide")

# الاتصال بقاعدة البيانات
engine = init_db()
from sqlalchemy.orm import sessionmaker
SessionLocal = sessionmaker(bind=engine)

st.title("📥 إدارة البريد الوارد")

# إنشاء تبويبات
tab1, tab2, tab3 = st.tabs(["📝 تسجيل جديد", "📋 قائمة البريد الوارد", "🔍 بحث متقدم"])

with tab1:
    st.subheader("تسجيل بريد وارد جديد")
    
    with st.form("entrant_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            reference = st.text_input("رقم المرجع*")
            date_reception = st.date_input("تاريخ الاستلام*", datetime.now())
            type_document = st.selectbox(
                "نوع الوثيقة",
                ["مراسلة", "مذكرة", "تقرير", "قرار", "تعميم", "آخر"]
            )
            
        with col2:
            # جلب الجهات من قاعدة البيانات
            db = SessionLocal()
            contacts = get_contacts(db)
            db.close()
            
            contact_names = [c.nom for c in contacts]
            expediteur_nom = st.selectbox("المرسل", [""] + contact_names)
            
            if expediteur_nom == "":
                expediteur_nom = st.text_input("أدخل اسم المرسل (إذا غير موجود)")
            
            priorite = st.selectbox("الأولوية", ["عادية", "عاجلة", "مهمة"])
            statut = st.selectbox("الحالة", ["غير معالج", "معالج", "متابعة"])
        
        objet = st.text_area("الموضوع*", height=100)
        fichier_joint = st.file_uploader("رفع ملف مرفق", type=['pdf', 'doc', 'docx', 'jpg', 'png'])
        notes = st.text_area("ملاحظات إضافية")
        
        submitted = st.form_submit_button("💾 حفظ البريد الوارد")
        
        if submitted:
            if not reference or not objet:
                st.error("الرجاء ملء جميع الحقول الإلزامية (*)")
            else:
                # حفظ الملف المرفق
                fichier_path = None
                if fichier_joint:
                    upload_dir = "uploads/entrant"
                    os.makedirs(upload_dir, exist_ok=True)
                    fichier_path = f"{upload_dir}/{reference}_{fichier_joint.name}"
                    with open(fichier_path, "wb") as f:
                        f.write(fichier_joint.getbuffer())
                
                # إيجاد أو إنشاء جهة المرسل
                db = SessionLocal()
                expediteur_id = None
                
                if expediteur_nom and expediteur_nom != "":
                    # البحث عن الجهة في قاعدة البيانات
                    existing_contact = db.query(Contact).filter(Contact.nom == expediteur_nom).first()
                    
                    if existing_contact:
                        expediteur_id = existing_contact.id
                    else:
                        # إنشاء جهة جديدة
                        new_contact = Contact(
                            nom=expediteur_nom,
                            type_contact="غير محدد"
                        )
                        db.add(new_contact)
                        db.commit()
                        db.refresh(new_contact)
                        expediteur_id = new_contact.id
                
                # حفظ في قاعدة البيانات
                entrant_data = {
                    "reference": reference,
                    "date_reception": date_reception,
                    "type_document": type_document,
                    "expediteur_id": expediteur_id,
                    "objet": objet,
                    "fichier_joint": fichier_path,
                    "priorite": priorite,
                    "statut": statut,
                    "notes": notes
                }
                
                try:
                    # استخدام create_entrant من CRUD
                    from database.crud import create_entrant
                    create_entrant(db, **entrant_data)
                    st.success("✅ تم حفظ البريد الوارد بنجاح!")
                    
                    # عرض زر للطباعة
                    if st.button("🖨️ طباعة بريدريو الاستلام"):
                        from docx import Document
                        
                        doc = Document()
                        doc.add_heading('برذريو استلام', 0)
                        
                        doc.add_paragraph(f'رقم المرجع: {reference}')
                        doc.add_paragraph(f'تاريخ الاستلام: {date_reception}')
                        
                        if expediteur_nom:
                            doc.add_paragraph(f'المرسل: {expediteur_nom}')
                        
                        doc.add_paragraph(f'الموضوع: {objet}')
                        doc.add_paragraph(f'نوع الوثيقة: {type_document}')
                        doc.add_paragraph(f'الأولوية: {priorite}')
                        
                        doc.save(f'bordereau_{reference}.docx')
                        st.success("تم إنشاء البرذريو بنجاح")
                        
                except Exception as e:
                    st.error(f"حدث خطأ: {str(e)}")
                    import traceback
                    st.error(traceback.format_exc())
                finally:
                    db.close()

with tab2:
    st.subheader("قائمة البريد الوارد")
    
    db = SessionLocal()
    entrants = db.query(Entrant).all()
    
    if entrants:
        data = []
        for e in entrants:
            # جلب اسم المرسل إذا كان موجودًا
            expediteur_name = "غير محدد"
            if e.expediteur_id:
                contact = db.query(Contact).filter(Contact.id == e.expediteur_id).first()
                if contact:
                    expediteur_name = contact.nom
            
            data.append({
                "المرجع": e.reference,
                "تاريخ الاستلام": e.date_reception,
                "المرسل": expediteur_name,
                "الموضوع": e.objet[:50] + "..." if len(e.objet) > 50 else e.objet,
                "النوع": e.type_document,
                "الأولوية": e.priorite,
                "الحالة": e.statut,
                "ملاحظات": e.notes[:30] + "..." if e.notes and len(e.notes) > 30 else e.notes
            })
        
        df = pd.DataFrame(data)
        st.dataframe(df, use_container_width=True)
        
        # خيارات التصفية
        col1, col2, col3 = st.columns(3)
        with col1:
            filter_statut = st.selectbox("تصفية حسب الحالة", ["الكل"] + list(df["الحالة"].unique()))
        with col2:
            filter_priorite = st.selectbox("تصفية حسب الأولوية", ["الكل"] + list(df["الأولوية"].unique()))
        
        if filter_statut != "الكل":
            df = df[df["الحالة"] == filter_statut]
        if filter_priorite != "الكل":
            df = df[df["الأولوية"] == filter_priorite]
            
        st.dataframe(df, use_container_width=True)
    else:
        st.info("لا توجد سجلات للبريد الوارد")
    
    db.close()

with tab3:
    st.subheader("بحث متقدم في البريد الوارد")
    
    search_col1, search_col2 = st.columns(2)
    
    with search_col1:
        search_ref = st.text_input("بحث برقم المرجع")
        search_sender = st.text_input("بحث بالمرسل")
        
    with search_col2:
        search_subject = st.text_input("بحث بالموضوع")
        date_range = st.date_input("نطاق التاريخ", [])
    
    if st.button("🔍 بدء البحث"):
        db = SessionLocal()
        query = db.query(Entrant)
        
        if search_ref:
            query = query.filter(Entrant.reference.contains(search_ref))
        if search_subject:
            query = query.filter(Entrant.objet.contains(search_subject))
        
        entrants = query.all()
        
        if entrants:
            # إذا كان هناك بحث عن المرسل، فلنبحث في جدول الجهات أولاً
            if search_sender:
                contacts = db.query(Contact).filter(Contact.nom.contains(search_sender)).all()
                contact_ids = [c.id for c in contacts]
                entrants = [e for e in entrants if e.expediteur_id in contact_ids]
        
        db.close()
        
        if entrants:
            data = []
            for e in entrants:
                # جلب اسم المرسل
                expediteur_name = "غير محدد"
                if e.expediteur_id:
                    db_temp = SessionLocal()
                    contact = db_temp.query(Contact).filter(Contact.id == e.expediteur_id).first()
                    if contact:
                        expediteur_name = contact.nom
                    db_temp.close()
                
                data.append({
                    "المرجع": e.reference,
                    "تاريخ الاستلام": e.date_reception,
                    "المرسل": expediteur_name,
                    "الموضوع": e.objet,
                    "الحالة": e.statut
                })
            st.dataframe(pd.DataFrame(data))
        else:
            st.warning("لم يتم العثور على نتائج")